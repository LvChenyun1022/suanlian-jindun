"""用模拟数据填写官方合同示范文本空白条款，生成 *_filled.docx（随后由 Word COM 转 PDF）。

合规约定：
- 所有填写值均为合成"示例"系企业名/合成证件样式，不含真实企业或个人信息；
- 产物仅用于外部效度测试，标注为"官方示范文本+模拟填写"，不得称为真实合同；
- 输出文件位于 data/external/（.gitignore 覆盖，不入库不上传）。
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
EXT = ROOT / "data" / "external"


def replace_in_paragraph(p, old: str, new: str) -> bool:
    """段落级替换：合并 run 文本替换后写回第一个 run（填空段落格式统一，可接受）。"""
    full = "".join(r.text for r in p.runs)
    if old not in full:
        return False
    full = full.replace(old, new, 1)
    if not p.runs:
        p.add_run(full)
        return True
    p.runs[0].text = full
    for r in p.runs[1:]:
        r.text = ""
    return True


def fill_industrial_sale() -> Path:
    src = EXT / "GF2000_industrial_sale_template.docx"
    dst = EXT / "GF2000_industrial_sale_filled.docx"
    d = Document(src)

    repl = [
        ("合同编号：", "合同编号：SL-GM-2026-0101"),
        ("出卖人：", "出卖人：示例算力设备（深圳）有限公司"),
        ("签订地点：", "签订地点：上海市浦东新区"),
        ("买受人：", "买受人：示例智算科技（上海）有限公司"),
        ("签订时间：", "签订时间：2026 年 3 月 15 日"),
        ("结算方式、时间及地点：",
         "结算方式、时间及地点：货到验收合格后 90 日内以电汇方式支付至出卖人指定账户"),
    ]
    missed = []
    for old, new in repl:
        if not any(replace_in_paragraph(p, old, new) for p in d.paragraphs):
            missed.append(old)

    # 标的第一条表格：第 2 行（索引 2）填 GPU 设备一行
    tb = d.tables[0]
    row = tb.rows[2]
    vals = ["GPU 服务器", "示例智造", "GPU-A800-80G", "示例智造（东莞）有限公司",
            "台", "64", "125000.00", "8000000.00"]
    seen = []
    for ci, cell in enumerate(row.cells):
        if id(cell._tc) in seen:
            continue
        seen.append(id(cell._tc))
        idx = len(seen) - 1
        if idx < len(vals):
            cell.paragraphs[0].add_run(vals[idx])
    # 合计大写行
    for cell in tb.rows[6].cells:
        p = cell.paragraphs[0]
        if "合计人民币金额" in p.text:
            p.add_run("捌佰万元整（￥8000000.00）")
            break

    d.save(dst)
    return dst, missed


def fill_data_service() -> Path:
    src = EXT / "GF-2025-2616_data_service_template.docx"
    dst = EXT / "GF-2025-2616_data_service_filled.docx"
    d = Document(src)

    repl = [
        ("甲方（委托方）：", "甲方（委托方）：示例数据科技（上海）有限公司"),
        ("乙方（受托方）：", "乙方（受托方）：示例云算服务（北京）有限公司"),
        ("数据名称：                                         。",
         "数据名称：GPU 算力集群运行日志数据集。"),
        ("处理期限、处理环境等要求",
         "处理期限自 2026 年 7 月 1 日至 2027 年 6 月 30 日、处理环境等要求"),
        ("费用总额为      （大写：                                                  ）。",
         "费用总额为 1,200,000.00 元（大写：壹佰贰拾万元整）。"),
        ("有效期至    年    月    日止", "有效期至 2027 年 6 月 30 日止"),
    ]
    missed = []
    for old, new in repl:
        if not any(replace_in_paragraph(p, old, new) for p in d.paragraphs):
            missed.append(old)
    d.save(dst)
    return dst, missed


def main() -> int:
    all_missed = {}
    for fn in (fill_industrial_sale, fill_data_service):
        dst, missed = fn()
        print("written:", dst.name)
        if missed:
            all_missed[dst.name] = missed
    if all_missed:
        print("MISSED replacements:", all_missed)
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
